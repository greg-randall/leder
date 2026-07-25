"""Stage E: Generate a .docx file with Word comments from verified claims.

Each footnote marker in the sourced article becomes a Word comment
anchored to the surrounding text. Upload to Google Drive → open with
Google Docs → all comments appear in the sidebar for collaboration.

Requires python-docx >= 1.2.0.
"""
from __future__ import annotations

import json
import re
import sys
from docx import Document
from docx.shared import Pt


def convert(article_sourced_md: str, findings_json: str, output_docx: str) -> None:
    with open(findings_json, encoding="utf-8") as f:
        claims_data = json.load(f)
    claims_by_id = {}
    items = claims_data.get("findings", [])
    for c in items:
        cid = c.get("finding_id") or c.get("claim_id", "")
        claims_by_id[cid] = c

    with open(article_sourced_md, encoding="utf-8") as f:
        md_text = f.read()

    # Split off sources section
    body = md_text.split("\n---\n\n## Sources\n")[0]
    body = re.sub(r'^# ⚠️ UNPLACED CLAIMS.*?\n---\n\n', '', body, flags=re.DOTALL)

    doc = Document()

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11.5)

    fn_pattern = re.compile(r'\[\^(\d+)\]')
    comment_count = 0

    for para_text in body.split('\n\n'):
        para_text = para_text.strip()
        if not para_text:
            continue

        # Headings
        if para_text.startswith('# '):
            doc.add_heading(para_text[2:].strip(), level=1)
            continue
        if para_text.startswith('## '):
            doc.add_heading(para_text[3:].strip(), level=2)
            continue
        if para_text.startswith('### '):
            doc.add_heading(para_text[4:].strip(), level=3)
            continue

        # Body paragraph — split on footnote markers to create targeted runs
        parts = fn_pattern.split(para_text)
        if len(parts) == 1:
            doc.add_paragraph(_clean(parts[0]))
            continue

        p = doc.add_paragraph()
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text
                p.add_run(_clean(part))
            else:
                # Footnote marker — build comment anchored to preceding text
                fn_id = f"c{int(part):04d}"
                claim = claims_by_id.get(fn_id)
                comment_count += 1

                if claim:
                    severity = str(claim.get("severity", "WARNING")).upper()
                    if severity == "PASS":
                        v_symbol = "✓"
                    elif severity == "CRITICAL":
                        v_symbol = "✗"
                    else:
                        v_symbol = "?"

                    text = (
                        f"{v_symbol} [{claim.get('check_type', 'check')}]\n"
                        f"Claim: {claim.get('claim_text', '')}\n"
                        f"Rationale: {claim.get('agent_summary') or claim.get('rationale', '')}"
                    )
                    excerpt = claim.get("source_excerpt")
                    if excerpt:
                        text += f'\nSource text: "{excerpt}"'
                    sp = claim.get("source_path")
                    su = claim.get("source_url")
                    if sp and sp != "null":
                        text += f"\nSource: {sp}"
                    if su and su != "null":
                        text += f"\nURL: {su}"
                    rec = claim.get("recommended_action")
                    if rec:
                        text += f"\n---\nRecommendation: {rec}"
                else:
                    text = "[No verification data]"
                    print(f"  stage-e: no claims.json entry for footnote [^{part}] "
                          f"(looked up as {fn_id}) — comment will show placeholder text",
                          file=sys.stderr)

                # Anchor to previous run(s) — the text right before this marker
                if p.runs:
                    doc.add_comment(
                        runs=[p.runs[-1]],
                        text=text,
                        author="Fact Check",
                        initials="FC",
                    )

    # Insert a superscript comment reference number at each anchor
    # (python-docx doesn't add visible markers automatically, so we don't either —
    # the comments appear in the sidebar without inline markers)

    # Save and clean up — flatten "Default Paragraph Font" runs
    doc.save(output_docx)

    passed = sum(1 for c in claims_by_id.values() if str(c.get("severity", "")).upper() == "PASS")
    criticals = sum(1 for c in claims_by_id.values() if str(c.get("severity", "")).upper() == "CRITICAL")
    warnings = sum(1 for c in claims_by_id.values() if str(c.get("severity", "")).upper() == "WARNING")

    print(f"DOCX written → {output_docx}")
    print(f"  {comment_count} comments ({passed}✓/{criticals}✗/{warnings}?)")
    print("  Upload to Google Drive → open with Google Docs → comments in sidebar")


def _clean(text: str) -> str:
    """Strip basic inline markdown."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def run_stage_e(article_sourced_md: str, findings_json: str, output_docx: str = "") -> str:
    if not output_docx:
        output_docx = article_sourced_md.replace(".md", ".docx")
    convert(article_sourced_md, findings_json, output_docx)
    return output_docx
