"""Tests for Stage E: .docx generation with Word comments.

Note: the "\\n---\\n\\n## Sources\\n" block and everything after it in the
markdown fixtures below is discarded by ``convert()`` (it splits the body off
at that exact delimiter) -- the hand-written "[^1]: **[...]**..." footnote
text is never read by the code under test. It's included only for fixture
readability; editing it has no effect on comment output. What actually
drives each comment's content is the corresponding entry in ``claims.json``.
"""
from __future__ import annotations

import json

from docx import Document

from pipeline.stage_e_docx import run_stage_e


def _write_claims_json(path, findings):
    path.write_text(json.dumps({"findings": findings}))


def test_happy_path_creates_comment(tmp_path):
    md = tmp_path / "sourced.md"
    md.write_text(
        "# Test Article\n\n"
        "The facility was fined in 2023.[^1]\n\n"
        "\n---\n\n## Sources\n"
        "[^1]: **[✓ fact_check]** [Original] The facility was fined. — Confirmed by EPA records.\n"
        "    Source: `epa/records.md`\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [{
        "finding_id": "c0001", "check_type": "fact_check", "severity": "PASS",
        "claim_text": "The facility was fined.", "agent_summary": "Confirmed by EPA records.",
        "source_path": "epa/records.md",
    }])

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    doc = Document(out)

    assert doc.paragraphs[0].text == "Test Article"
    comments = list(doc.comments)
    assert len(comments) == 1
    assert "fact_check" in comments[0].text
    assert "Confirmed by EPA records." in comments[0].text
    assert "epa/records.md" in comments[0].text
    assert comments[0].text.startswith("✓")
    assert comments[0].author == "Fact Check"


def test_multiple_findings_produce_multiple_comments(tmp_path):
    md = tmp_path / "sourced.md"
    md.write_text(
        "# Test Article\n\n"
        "First claim here.[^1] Second claim here.[^2]\n\n"
        "\n---\n\n## Sources\n"
        "[^1]: **[✓ fact_check]** [Original] First claim. — Rationale one.\n"
        "    Source: `a.md`\n"
        "[^2]: **[✗ fact_check]** [Original] Second claim. — Rationale two.\n"
        "    Source: `b.md`\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [
        {"finding_id": "c0001", "check_type": "fact_check", "severity": "PASS",
         "claim_text": "First claim.", "agent_summary": "Rationale one.", "source_path": "a.md"},
        {"finding_id": "c0002", "check_type": "fact_check", "severity": "CRITICAL",
         "claim_text": "Second claim.", "agent_summary": "Rationale two.", "source_path": "b.md"},
    ])

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    comments = list(Document(out).comments)
    assert len(comments) == 2
    texts = [c.text for c in comments]
    # Pair each rationale with its own claim's severity symbol and claim text,
    # so a mix-up between c0001's and c0002's data (e.g. footnote [^1] picking
    # up c0002's fields) would fail this test.
    assert any("Rationale one." in t and t.startswith("✓") and "First claim." in t for t in texts)
    assert any("Rationale two." in t and t.startswith("✗") and "Second claim." in t for t in texts)


def test_severity_symbol_mapping(tmp_path):
    md = tmp_path / "sourced.md"
    md.write_text(
        "# Test\n\nClaim A.[^1] Claim B.[^2] Claim C.[^3]\n\n"
        "\n---\n\n## Sources\n"
        "[^1]: **[✓ fact_check]** [Original] Claim A. — r.\n    Source: `a.md`\n"
        "[^2]: **[✗ fact_check]** [Original] Claim B. — r.\n    Source: `b.md`\n"
        "[^3]: **[? fact_check]** [Original] Claim C. — r.\n    Source: `c.md`\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [
        {"finding_id": "c0001", "check_type": "fact_check", "severity": "PASS",
         "claim_text": "Claim A.", "agent_summary": "r", "source_path": "a.md"},
        {"finding_id": "c0002", "check_type": "fact_check", "severity": "CRITICAL",
         "claim_text": "Claim B.", "agent_summary": "r", "source_path": "b.md"},
        {"finding_id": "c0003", "check_type": "fact_check", "severity": "WARNING",
         "claim_text": "Claim C.", "agent_summary": "r", "source_path": "c.md"},
    ])

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    comments = list(Document(out).comments)
    assert len(comments) == 3
    texts = [c.text for c in comments]
    # Pair each symbol with the specific claim it must belong to, so a mutation
    # that swaps the severity->symbol branches (e.g. PASS->✗, CRITICAL->✓) or
    # that misattributes a comment to the wrong footnote would fail this test.
    assert any(t.startswith("✓") and "Claim A" in t for t in texts)
    assert any(t.startswith("✗") and "Claim B" in t for t in texts)
    assert any(t.startswith("?") and "Claim C" in t for t in texts)


def test_headings_render_as_word_headings(tmp_path):
    md = tmp_path / "sourced.md"
    md.write_text(
        "# Main Title\n\n## Section One\n\n### Subsection\n\nBody text.\n\n"
        "\n---\n\n## Sources\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [])

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    doc = Document(out)
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Heading 3" in styles


def test_malformed_claims_json_missing_fields_does_not_crash(tmp_path):
    """Findings with missing optional fields still produce a comment
    instead of crashing."""
    md = tmp_path / "sourced.md"
    md.write_text(
        "# Test\n\nA sparse claim.[^1]\n\n"
        "\n---\n\n## Sources\n"
        "[^1]: **[✓ fact_check]** [Original] A sparse claim. — r.\n    Source: `none`\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [{"finding_id": "c0001"}])  # only the ID, nothing else

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    comments = list(Document(out).comments)
    assert len(comments) == 1
    assert "check" in comments[0].text  # falls back to the default check_type label


def test_unmatched_footnote_id_gets_placeholder_and_warns(tmp_path, capsys):
    """A footnote marker whose ID isn't in claims.json must not silently
    vanish with zero signal -- drives the warning fix in this task."""
    md = tmp_path / "sourced.md"
    md.write_text(
        "# Test\n\nAn orphaned marker.[^1]\n\n"
        "\n---\n\n## Sources\n"
        "[^1]: **[✓ fact_check]** [Original] An orphaned marker. — r.\n    Source: `none`\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [])  # no c0001 entry at all

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    comments = list(Document(out).comments)
    assert len(comments) == 1
    assert comments[0].text == "[No verification data]"
    assert "c0001" in capsys.readouterr().err


def test_unplaced_claims_block_is_fully_stripped(tmp_path):
    """Regression test: stage_c_rebuild's unplaced-claims block (heading +
    intro + claim bullets) must not leak into the docx as visible body
    paragraphs. Uses real build_unplaced_warning() output assembled the
    same way run_stage_c actually joins it, so this can't drift out of
    sync with the real producer the way the original bug did (the strip
    regex only matched through the first blank line, leaving the intro
    paragraph and claim bullets in the visible body)."""
    from pipeline.stage_c_rebuild import build_unplaced_warning
    from types import SimpleNamespace
    from pipeline.models import Verdict

    unplaced_claim = SimpleNamespace(
        claim_id="c099", claim_text="An unplaceable claim.", verdict=Verdict.SUPPORTED,
        source_path="orphan.md", source_url=None, rationale="Could not locate in article.",
        source_quote="a quote that does not appear anywhere",
    )
    unplaced_block = build_unplaced_warning([unplaced_claim])

    md = tmp_path / "sourced.md"
    md.write_text(
        unplaced_block + "\n---\n\n"
        "# Test Article\n\nOrdinary article text.\n\n"
        "\n---\n\n## Sources\n"
    )
    claims = tmp_path / "claims.json"
    _write_claims_json(claims, [])

    out = run_stage_e(str(md), str(claims), str(tmp_path / "out.docx"))
    doc = Document(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)

    assert "UNPLACED CLAIMS" not in body_text
    assert "MANUAL REVIEW REQUIRED" not in body_text
    assert "An unplaceable claim." not in body_text
    assert "Ordinary article text." in body_text
